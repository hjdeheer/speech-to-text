import datetime
# Create an folder inside audio_out_path to save the files based on the date, h m and seconds
import os
from dotenv import load_dotenv
import humanfriendly
import matplotlib.pyplot as plt
import numpy as np
import torch
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from pyannote.audio.pipelines.utils.hook import ProgressHook
from pyannote.core import Annotation
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline

generate_kwargs_asr = {

}

generate_kwargs_diarization = {
    # 'num_speakers': 2,
    # 'max_speakers': 5,
}

audio_in_path = "in/"
audio_out_path = "out/"
return_timestamps = True

save_type = "docx_pdf"

load_dotenv()

def load_audio(audio_path_folder: str) -> tuple[list[dict], list[str]]:
    """
    Load audio file from path using librosa.
    :param audio_path_folder: path to audio file
    :return: dictionary with keys "path", "sampling_rate", "array"
    """
    import librosa

    # Read all files inside the audio path folder using os listdir
    audio_samples = []
    names = []
    import os

    for file in os.listdir(audio_path_folder):
        audio_path = os.path.join(audio_path_folder, file)
        audio, sampling_rate = librosa.load(audio_path, sr=None)
        name = audio_path.split("/")[-1]
        sample = {"path": name, "sampling_rate": sampling_rate, "array": audio}
        audio_samples.append(sample)
        names.append(name)

    return audio_samples, names


def write_to_file(audio_names: list[str], results: list[dict], speaker_annotations: list[Annotation], audio_in_path: str, audio_out_path: str, save_type: str, return_timestamps: bool = False) -> None:
    """
    Write audio transcription to file.
    :param audio_names: Original audio names
    :param results: Results from ASR
    :param audio_out_path: Path to save audio files
    :param save_type: Type of file to save to (docx, pdf)
    :return: None
    """
    date = datetime.datetime.now().strftime("%Y-%m-%d-%H:%M-%S")
    audio_out_path = os.path.join(audio_out_path, date)
    os.makedirs(audio_out_path, exist_ok=True)

    # For each audio file, create a docx file with the transcription
    for i, (sample, result, annotations) in enumerate(zip(audio_names, results, speaker_annotations)):

        # Copy curr sample from audio in path to audio out path
        import shutil
        shutil.copy(os.path.join(audio_in_path, sample), audio_out_path)
        # Write to file based on the save_type
        for doc_type in save_type.split("_"):
            if doc_type == "docx":
                doc = Document()

                # Set font size and style
                doc.styles['Normal'].font.size = Pt(12)
                doc.styles['Normal'].font.name = 'Arial'

                doc.add_heading(f"Transcription for {sample} - {date}", level=1)
                # Add a line break
                doc.add_paragraph("")
                doc.add_paragraph(result["text"])
                doc.save(os.path.join(audio_out_path, f"transcription_{sample.split('.')[0]}.docx"))
                if return_timestamps:
                    write_timestamps_to_docx(sample, result, annotations, date, audio_out_path)
            elif doc_type == "pdf":
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(200, 10, f"Transcription for {sample} - {date}", ln=True)
                # Add a line break
                pdf.ln(10)
                # Add the long text using multi_cell
                pdf.multi_cell(0, 10, result["text"])
                pdf.output(os.path.join(audio_out_path, f"transcription_{sample.split('.')[0]}.pdf"))
                if return_timestamps:
                    write_timestamps_to_pdf(sample, result, date, audio_out_path)
            else:
                raise ValueError(f"Invalid save_type {doc_type}")


def process_diarization(annotation: Annotation):
    """
    Process the diarization annotation to get the segments and speakers.
    :param annotation: Annotation object from pyannote
    :return: segments, speakers
    """
    segments = []
    speakers = []
    for segment, speaker in annotation._tracks.items():
        segments.append(segment)
        speakers.append(list(speaker.values())[0])
    return segments, speakers


def generate_n_colors(speakers):
    # Create a colormap from 'tab20', a business-friendly color palette

    n = len(speakers)
    cmap = plt.get_cmap('tab10')

    # If n is greater than the number of distinct colors in 'tab20', we will interpolate new colors
    colors = [cmap(i) for i in np.linspace(0, 1, n)]

    # Return the colors as RGBColor objects
    colors = [RGBColor(int(c[0] * 255), int(c[1] * 255), int(c[2] * 255)) for c in colors]

    color_names = [closest_color_name((int(c[0]), int(c[1]), int(c[2]))) for c in colors]

    speaker_color_map = dict(zip(speakers, colors))

    return speaker_color_map, color_names


import webcolors


def closest_color_name(rgb_tuple):
    try:
        # Get the exact name if possible
        return webcolors.rgb_to_name(rgb_tuple)
    except ValueError:
        # If exact name is not found, find the closest color
        closest_name = None
        min_diff = float('inf')
        for name in webcolors.names("css3"):
            r, g, b = webcolors.name_to_rgb(name)
            diff = ((r - rgb_tuple[0]) ** 2 +
                    (g - rgb_tuple[1]) ** 2 +
                    (b - rgb_tuple[2]) ** 2)
            if diff < min_diff:
                min_diff = diff
                closest_name = name
        return closest_name


def write_timestamps_to_docx(sample, result, annotations, date, write_path):
    segments, speakers = process_diarization(annotations)

    # Get n unique speakers
    n_speakers = set(speakers)

    # Get n unique color and create map of speaker to map
    color_map, color_names = generate_n_colors(n_speakers)

    # Create a new document
    doc = Document()
    # Set font size and style
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].font.name = 'Arial'

    doc.add_heading(f"Transcription with timestamps for {sample} - {date}\n", level=1)
    doc.add_paragraph(f"{generate_description(result, segments, speakers)}\n", style='Heading 3')

    # Introduce speakers with appropriate colors
    for (speaker, color), name in zip(color_map.items(), color_names):
        # Apply the speaker color
        doc.add_paragraph(f"{speaker}: {name}", style='Heading 3').runs[0].font.color.rgb = color

    # Add a line break
    doc.add_paragraph("")

    for chunk in result["chunks"]:
        curr_text = chunk["text"]
        times = chunk["timestamp"]
        start_time = str(times[0]) + " s"
        end_time = str(times[1]) + " s"
        middle_s = (chunk["timestamp"][0] + chunk["timestamp"][1]) / 2
        speaker = get_speaker_of_chunk(middle_s, annotations)
        if speaker is not None:
            color = color_map[speaker]
            doc.add_paragraph(f"{start_time} - {end_time}: {curr_text}").runs[0].font.color.rgb = color
        else:
            doc.add_paragraph(f"{start_time} - {end_time}: {curr_text}")
    doc.save(os.path.join(write_path, f"transcription_timestamped_{sample.split('.')[0]}.docx"))


def get_speaker_of_chunk(middle_s, annotations):
    for segment, speaker in annotations._tracks.items():
        # Create new segment with 1second before and after
        if segment.overlaps(middle_s):
            return list(speaker.values())[0]

    # Ok we have not found a speaker. Maybe this current chunk is a pause, lets see if the left segment and the right segment have the same speaker
    for (segment_first, speaker_first), (segment_second, speaker_second) in zip(list(annotations._tracks.items()), list(annotations._tracks.items())[1:]):
        # If the middle_s is between the end of the first segment and the start of the second segment
        if segment_first.end <= middle_s <= segment_second.start:
            # If the speakers are the same, return the speaker
            if speaker_first == speaker_second:
                return list(speaker_first.values())[0]
            # If the speakers are different, return the closest speaker
            else:
                if middle_s - segment_first.end < segment_second.start - middle_s:
                    return list(speaker_first.values())[0]
                else:
                    return list(speaker_second.values())[0]
    return None


def generate_description(result, segments, speakers):
    length = humanfriendly.format_timespan(segments[-1].end)
    num_words = sum([len(chunk["text"].split()) for chunk in result["chunks"]])
    n_speakers = len(set(speakers))
    description = f"Transcription is {length} long, and contains {num_words} words spoken by {n_speakers} different speakers."
    return description


def write_timestamps_to_pdf(sample, result, date, write_path):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, f"Transcription with timestamps for {sample} - {date}", ln=True)
    for chunk in result["chunks"]:
        curr_text = chunk["text"]
        times = chunk["timestamp"]
        start_time = str(times[0]) + " s"
        end_time = str(times[1]) + " s"
        pdf.multi_cell(0, 10, f"{start_time} - {end_time}: {curr_text}")
    pdf.output(os.path.join(write_path, f"transcription_timestamped_{sample.split('.')[0]}.pdf"))


def asr(sample: dict, return_timestamps: bool, generate_kwargs: dict = None) -> dict:
    """
    Automatic Speech Recognition (ASR) using Hugging Face's pipeline.
    :param audio_samples: should be a list of dictionaries with keywords "path" and "sampling_rate" and "array"
    :param generate_kwargs:
    :return:
    """
    device = "cuda:0" if torch.cuda.is_available() else "cpu"
    torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32

    model_id = "openai/whisper-large-v3"

    model = AutoModelForSpeechSeq2Seq.from_pretrained(
        model_id, torch_dtype=torch_dtype, low_cpu_mem_usage=True, use_safetensors=True
    )
    model.to(device)

    processor = AutoProcessor.from_pretrained(model_id)

    pipe = pipeline(
        "automatic-speech-recognition",
        model=model,
        tokenizer=processor.tokenizer,
        feature_extractor=processor.feature_extractor,
        chunk_length_s=30,
        batch_size=64,
        torch_dtype=torch_dtype,
        device=device,
    )

    result = pipe(sample, return_timestamps=return_timestamps, generate_kwargs=generate_kwargs)
    return result


def diarization(audio_folder: str, **kwargs) -> list[Annotation]:
    """
    Speaker Diarization using pyannote.audio
    :param audio_folder: audio folder path
    :param kwargs: keyword arguments for the diarization pipeline
    :return: A list of annotations with annotations for each file.
    """
    # instantiate the pipeline

    from pyannote.audio import Pipeline
    pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1",
        use_auth_token=os.getenv("HUGGINGFACE_TOKEN"))

    pipeline.to(torch.device("cuda:0" if torch.cuda.is_available() else "cpu"))
    annotations = []
    # process each file in the audio folder
    for file in os.listdir(audio_folder):
        audio_path = os.path.join(audio_folder, file)
        # run the pipeline on an audio file
        with ProgressHook() as hook:
            annotation = pipeline(audio_path, hook=hook, **kwargs)

        annotations.append(annotation)
    return annotations


if __name__ == "__main__":
    # Start timer
    start = datetime.datetime.now()

    print("Loading audio...")
    audio_samples, names = load_audio(audio_in_path)

    print("Running ASR...")
    results = []

    for sample in audio_samples:
        result = asr(sample, return_timestamps=return_timestamps, generate_kwargs=generate_kwargs_asr)
        results.append(result)

    if return_timestamps:
        print("Running Speech Diarization...")
        speaker_annotations = diarization(audio_in_path)

    write_to_file(names, results, speaker_annotations, audio_in_path, audio_out_path, save_type, return_timestamps=return_timestamps)

    # End timer in seconds
    end = datetime.datetime.now()
    print(f"Time taken (s): {end - start}")
    print("Done!")
